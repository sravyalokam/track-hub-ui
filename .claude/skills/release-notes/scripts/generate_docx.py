#!/usr/bin/env python
"""
Render a Chorus PMO deployment-notes doc (.docx) from a JSON spec, matching
the exact styling reverse-engineered from a real template
(PMO-UAT-Release-Doc(04-June-2026).docx): centered dark-blue title, bold
bright-blue section headers, 12pt paragraph spacing, native Word bullet
list ("List Bullet" style, justified).

Usage:
    python generate_docx.py spec.json output.docx

Spec shape (see scripts/example_docx_spec.json for a full example):
{
  "app": "PMO_Application",
  "version": "2026.07.20",
  "source_env": "QA",
  "target_env": "UAT",
  "repos": [
    {"name": "Cognine-PMO-UI", "release": "{RELEASE_NUMBER}"},
    {"name": "Cognine-PMO-Backend", "release": "{RELEASE_NUMBER}"}
  ],
  "repo_label": "PMO Backend Repos",
  "stories_heading_label": "Users Stories Targeted",
  "sprint_label": null,
  "stories": [
    {"ticket_id": "15268", "is_bug": true, "text": "Percentage completion rollup incorrect for Removed/Blocker statuses"}
  ],
  "deployed_heading": "PMO App: Deployed Items",
  "items": [
    {"heading": "Sprint Workdays Accuracy", "text": "The Timeline now generates ..."},
    {"text": "A bullet with no side heading is also fine — heading is optional."}
  ]
}

Rules encoded here (do not deviate without checking the skill's SKILL.md):
- Omit the "Users Stories Targeted" section entirely if "stories" is empty/absent.
- repos: list every repo that shipped; each becomes its own plain (non-bullet)
  line under the repo_label heading.
- items: each becomes one bullet. "heading" is optional — when present it's
  rendered as a bold crisp lead-in ("Heading: description"); omit it for a
  plain bullet.
"""
import json
import sys

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

TITLE_BLUE = RGBColor(0x0F, 0x47, 0x61)
HEADING_BLUE = RGBColor(0x0F, 0x9E, 0xD5)
BLACK = RGBColor(0x00, 0x00, 0x00)


def set_run(r, bold=None, color=None):
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    if bold is not None:
        r.bold = bold
    if color is not None:
        r.font.color.rgb = color
    return r


def spaced_para(doc, alignment=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    if alignment is not None:
        p.alignment = alignment
    return p


def build(spec):
    doc = Document()

    title = f"{spec['app']}_{spec['version']} Deployment Notes"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(title), color=TITLE_BLUE)

    p = doc.add_paragraph()
    set_run(p.add_run("OVERALL STEPS"), bold=True, color=HEADING_BLUE)

    source_env = spec["source_env"]
    target_env = spec["target_env"]
    p = spaced_para(doc)
    set_run(p.add_run("During the Day: "), bold=True, color=BLACK)
    set_run(
        p.add_run(
            f"All things needed for deployment like raising pull requests from "
            f"{source_env} to {target_env} for ({target_env} release) along with "
            f"pull request links should be provided here."
        ),
        color=BLACK,
    )

    p = spaced_para(doc)
    set_run(p.add_run("Next Day"), bold=True, color=HEADING_BLUE)
    set_run(
        p.add_run(": Once production is completed, delete local branch cleanup party!"),
        color=BLACK,
    )

    stories = spec.get("stories") or []
    if stories:
        p = spaced_para(doc)
        heading_text = spec.get("stories_heading_label", "Users Stories Targeted")
        if spec.get("sprint_label"):
            heading_text += f" In {spec['sprint_label']}"
        set_run(p.add_run(heading_text + ":"), bold=True, color=HEADING_BLUE)
        for s in stories:
            p = spaced_para(doc)
            prefix = f"Bug {s['ticket_id']}: " if s.get("is_bug") else f"{s['ticket_id']}: "
            set_run(p.add_run(prefix + s["text"]), color=BLACK)

    p = spaced_para(doc)
    set_run(p.add_run(spec["repo_label"] + " – Release Pipelines"), bold=True, color=HEADING_BLUE)
    for repo in spec["repos"]:
        p = spaced_para(doc)
        set_run(p.add_run(repo["name"]), bold=True, color=BLACK)
        set_run(p.add_run(f" – Release {repo['release']}"), color=BLACK)

    p = spaced_para(doc)
    set_run(p.add_run(spec["deployed_heading"]), bold=True, color=HEADING_BLUE)

    p = spaced_para(doc)
    set_run(p.add_run("The following are the updates as per this release:"), color=BLACK)

    for item in spec["items"]:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        if item.get("heading"):
            set_run(p.add_run(item["heading"] + ": "), bold=True, color=BLACK)
        set_run(p.add_run(item["text"]), color=BLACK)

    return doc


def main():
    if len(sys.argv) != 3:
        print("usage: python generate_docx.py spec.json output.docx", file=sys.stderr)
        sys.exit(1)

    spec_path, out_path = sys.argv[1], sys.argv[2]
    with open(spec_path, "r", encoding="utf-8") as f:
        spec = json.load(f)

    doc = build(spec)

    try:
        doc.save(out_path)
    except PermissionError:
        # Target is almost certainly open in Word right now. Never fail silently
        # or overwrite blindly once it becomes writable again — save alongside
        # it with a suffix and tell the caller so they can reconcile by hand.
        import os

        root, ext = os.path.splitext(out_path)
        fallback = f"{root}_new{ext}"
        doc.save(fallback)
        print(
            f"WARNING: {out_path} is locked (likely open in Word). "
            f"Saved to {fallback} instead.",
            file=sys.stderr,
        )
        print(fallback)
        return

    print(out_path)


if __name__ == "__main__":
    main()
